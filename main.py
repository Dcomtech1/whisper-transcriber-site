import os
import uuid
from datetime import datetime
from typing import Optional

from fastapi import FastAPI, UploadFile, Form, Request
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

import whisper
from docx import Document

app = FastAPI(title="Nova Transcribe")
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

DEFAULT_MODEL = os.getenv("WHISPER_MODEL", "base")
MODEL = whisper.load_model(DEFAULT_MODEL)

OUTPUT_DIR = "outputs"
os.makedirs(OUTPUT_DIR, exist_ok=True)

@app.get("/health")
async def health():
    return {"ok": True, "model": DEFAULT_MODEL}

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "default_model": DEFAULT_MODEL})

@app.post("/api/transcribe")
async def transcribe_api(
    file: UploadFile,
    model_size: str = Form(DEFAULT_MODEL),
    task: str = Form("transcribe"),
    temperature: float = Form(0.0),
    word_timestamps: Optional[bool] = Form(False)
):
    global MODEL

    # Save upload
    base_name = os.path.splitext(file.filename or "audio")[0]
    ext = os.path.splitext(file.filename or "")[1].lower() or ".wav"
    uid = uuid.uuid4().hex[:8]
    input_path = os.path.join(OUTPUT_DIR, f"{base_name}_{uid}{ext}")
    with open(input_path, "wb") as f:
        f.write(await file.read())

    # Reload model if different
    current_model_name = getattr(MODEL, 'name', DEFAULT_MODEL)
    if model_size and model_size != current_model_name:
        try:
            MODEL = whisper.load_model(model_size)
        except Exception as e:
            return JSONResponse(status_code=400, content={"error": f"Failed to load model '{model_size}': {e}"})

    # Transcribe
    try:
        result = MODEL.transcribe(
            input_path,
            task=task,
            temperature=temperature,
            word_timestamps=bool(word_timestamps),
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Transcription failed: {e}"})

    text = (result.get("text") or "").strip()
    segments = result.get("segments", []) or []
    language = result.get("language", None)

    # Approx duration from last segment end
    duration = None
    if segments:
        try:
            duration = float(segments[-1].get("end", 0.0))
        except Exception:
            duration = None

    # Build DOCX
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_name = f"{base_name}_{ts}.docx"
    docx_path = os.path.join(OUTPUT_DIR, docx_name)

    doc = Document()
    doc.add_heading("Transcription", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Source: ").bold = True
    meta.add_run(file.filename or "Uploaded file")
    meta.add_run("\nModel: ").bold = True
    meta.add_run(model_size)
    if language:
        meta.add_run("\nDetected language: ").bold = True
        meta.add_run(str(language))
    if duration:
        meta.add_run("\nDuration: ").bold = True
        meta.add_run(f"{duration:.1f}s")
    doc.add_paragraph("")
    if text:
        doc.add_paragraph(text)

    if segments:
        doc.add_page_break()
        doc.add_heading("Segments", level=2)
        for s in segments:
            start = s.get("start", 0.0)
            end = s.get("end", 0.0)
            stext = (s.get("text") or "").strip()
            doc.add_paragraph(f"[{start:.2f}s → {end:.2f}s] {stext}")

    doc.save(docx_path)

    return {
        "ok": True,
        "text": text,
        "docx_file": f"/download/{docx_name}",
        "model": model_size,
        "language": language,
        "duration_seconds": duration,
        "segments": segments,
        "filename": file.filename,
    }

@app.get("/download/{fname}")
async def download_file(fname: str):
    path = os.path.join(OUTPUT_DIR, fname)
    if not os.path.exists(path):
        return JSONResponse(status_code=404, content={"error": "File not found"})
    return FileResponse(path, filename=fname)
